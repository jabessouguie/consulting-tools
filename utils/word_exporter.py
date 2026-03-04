"""
WordExporter — Export de contenu structuré vers fichiers .docx (python-docx).

Fonctions :
    export_to_word(content, path, title) → str  # Chemin du .docx généré
"""

from pathlib import Path
from typing import Any, Dict, List, Optional


def export_to_word(
    content: Dict[str, Any],
    path: str,
    title: str = "",
) -> str:
    """
    Génère un fichier .docx structuré depuis un dictionnaire de contenu.

    Args:
        content: Dict avec les clés :
            - "title" (str) : titre principal (optionnel, surchargé par `title`)
            - "sections" (List[Dict]) : liste de sections avec :
                - "heading" (str) : titre de section
                - "body" (str) : contenu texte (markdown basique supporté : **gras**)
                - "level" (int, optionnel) : niveau de heading 1–3 (défaut : 1)
        path: Chemin de sortie du fichier .docx.
        title: Titre principal (surcharge content["title"]).

    Returns:
        Chemin absolu du fichier .docx créé.

    Raises:
        ImportError: si python-docx n'est pas installé.
        OSError: si l'écriture du fichier échoue.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError as exc:
        raise ImportError(
            "python-docx n'est pas installé. Exécutez : pip install python-docx"
        ) from exc

    doc = Document()

    # Titre principal
    main_title = title or content.get("title", "Document")
    heading = doc.add_heading(main_title, level=0)

    # Sections
    sections: List[Dict[str, Any]] = content.get("sections", [])
    for section in sections:
        heading_text = section.get("heading", "")
        body_text = section.get("body", "")
        level = int(section.get("level", 1))

        if heading_text:
            doc.add_heading(heading_text, level=min(level, 3))

        if body_text:
            _add_body_paragraph(doc, body_text)

    # Sauvegarder
    output_path = Path(path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return str(output_path.resolve())


def _add_body_paragraph(doc: Any, text: str) -> None:
    """
    Ajoute un paragraphe au document, avec support basique du markdown :
    - **texte** → gras
    - Lignes commençant par "- " → liste à puces

    Args:
        doc: Document python-docx.
        text: Texte à insérer (potentiellement multi-lignes).
    """
    lines = text.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("- ") or stripped.startswith("• "):
            item_text = stripped[2:]
            para = doc.add_paragraph(style="List Bullet")
            _add_runs_with_bold(para, item_text)
        else:
            para = doc.add_paragraph()
            _add_runs_with_bold(para, stripped)


def _add_runs_with_bold(para: Any, text: str) -> None:
    """
    Ajoute des runs dans un paragraphe en interprétant **texte** comme gras.

    Args:
        para: Paragraphe python-docx.
        text: Texte pouvant contenir **bold**.
    """
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        run.bold = (i % 2 == 1)
