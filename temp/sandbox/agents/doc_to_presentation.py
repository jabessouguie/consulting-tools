"""
Agent de conversion de documents en presentation
Prend un ou plusieurs documents (md, pdf, docx) + public cible + objectif
et genere une presentation PPTX structuree
"""
import os
import sys
import json
from typing import Dict, Any, Optional, List
from datetime import datetime
from pathlib import Path

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from dotenv import load_dotenv
load_dotenv(os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), '.env'))

from utils.llm_client import LLMClient


class DocToPresentationAgent:
    """Agent pour convertir des documents en presentation PPTX"""

    def __init__(self):
        self.llm = LLMClient(max_tokens=8192)
        self.base_dir = Path(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        self.consultant_info = {
            'name': os.getenv('CONSULTANT_NAME', 'Jean-Sebastien Abessouguie Bayiha'),
            'company': os.getenv('COMPANY_NAME', 'Consulting Tools'),
        }

    def parse_document(self, file_path: str, file_content: bytes = None, filename: str = "") -> str:
        """Parse un document et retourne son contenu texte"""
        ext = Path(filename or file_path).suffix.lower()

        if ext == '.md':
            if file_content:
                return file_content.decode('utf-8')
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()

        elif ext == '.pdf':
            from PyPDF2 import PdfReader
            import io
            if file_content:
                reader = PdfReader(io.BytesIO(file_content))
            else:
                reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text

        elif ext == '.docx':
            from docx import Document
            import io
            if file_content:
                doc = Document(io.BytesIO(file_content))
            else:
                doc = Document(file_path)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

        elif ext == '.txt':
            if file_content:
                return file_content.decode('utf-8')
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()

        return ""

    def analyze_and_structure(self, documents_text: str, target_audience: str, objective: str) -> List[Dict]:
        """Analyse les documents et genere la structure des slides en JSON"""
        print("  [2/4] Analyse et structuration des slides...")

        system_prompt = f"""Tu es un expert en creation de presentations professionnelles pour {self.consultant_info['company']}.
Tu crees des presentations claires, visuelles et percutantes adaptees au public cible.

REGLES :
- Maximum 1 idee par slide
- Titres courts et percutants
- Bullet points concis (max 5 par slide)
- Inclure des slides de type: cover, section, content, stat, highlight, diagram, closing
- Adapter le ton et le vocabulaire au public cible
- Structurer logiquement : intro, developpement par themes, conclusion"""

        prompt = f"""A partir des documents suivants, genere une presentation structuree.

PUBLIC CIBLE : {target_audience}
OBJECTIF : {objective}

CONTENU DES DOCUMENTS :
{documents_text[:8000]}

Genere la structure des slides au format JSON. Chaque slide doit avoir un type parmi:
- "cover" : slide de couverture (title, subtitle)
- "section" : separateur de section (title)
- "content" : contenu avec puces (title, bullets[])
- "stat" : slide statistique (stat_value, stat_label, context)
- "highlight" : points cles (title, key_points[])
- "diagram" : diagramme (title, diagram_type, elements[])
- "closing" : slide de cloture (title, subtitle)

Reponds UNIQUEMENT avec le JSON :
```json
{{
  "title": "Titre de la presentation",
  "slides": [
    {{"type": "cover", "title": "...", "subtitle": "..."}},
    {{"type": "section", "title": "..."}},
    {{"type": "content", "title": "...", "bullets": ["...", "..."]}},
    ...
  ]
}}
```"""

        response = self.llm.generate(prompt=prompt, system_prompt=system_prompt, temperature=0.5)

        try:
            if '```json' in response:
                json_str = response.split('```json')[1].split('```')[0].strip()
            elif '```' in response:
                json_str = response.split('```')[1].split('```')[0].strip()
            else:
                json_str = response.strip()
            data = json.loads(json_str)
            return data.get('slides', [])
        except json.JSONDecodeError as e:
            print(f"  Erreur JSON: {e}")
            return []

    def build_pptx(self, slides: List[Dict], title: str = "Presentation") -> Optional[str]:
        """Construit le fichier PPTX a partir des slides"""
        print("  [3/4] Construction du PPTX...")

        try:
            from utils.pptx_generator import build_proposal_pptx

            output_dir = self.base_dir / "output"
            output_dir.mkdir(exist_ok=True)

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = output_dir / f"presentation_{timestamp}.pptx"
            template_path = self.base_dir / "Consulting Tools_Template_Palette 2026.pptx"

            pptx_path = build_proposal_pptx(
                template_path=str(template_path),
                slides_data=slides,
                output_path=str(output_path),
                consultant_info=self.consultant_info,
            )

            return str(Path(pptx_path).relative_to(self.base_dir))
        except Exception as e:
            print(f"  Erreur PPTX: {e}")
            return None

    def generate_images(self, slides: List[Dict]) -> List[Dict]:
        """Genere des images pour les slides qui en beneficieraient"""
        print("  [4/4] Generation des illustrations...")
        try:
            from utils.image_generator import NanoBananaGenerator
            generator = NanoBananaGenerator()

            output_dir = self.base_dir / "output" / "images"
            output_dir.mkdir(parents=True, exist_ok=True)

            for i, slide in enumerate(slides):
                if slide.get('type') in ('cover', 'section', 'highlight'):
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    img_path = str(output_dir / f"slide_{i}_{timestamp}.png")
                    prompt = f"Corporate tech illustration for a presentation slide titled '{slide.get('title', '')}'. Style: minimalist, professional, cool blue and warm gold lighting, isometric view."
                    result = generator.generate_image(prompt, img_path)
                    if result:
                        slide['image_path'] = result
        except Exception as e:
            print(f"  Erreur images: {e}")

        return slides

    def run(self, documents: List[Dict[str, Any]], target_audience: str, objective: str) -> Dict[str, Any]:
        """
        Pipeline complet : parse -> analyse -> images -> PPTX

        Args:
            documents: Liste de dicts {"filename": str, "content": bytes ou str}
            target_audience: Public cible
            objective: Objectif de la presentation
        """
        print(f"\n  GENERATION DE PRESENTATION")
        print(f"  Public: {target_audience}")
        print(f"  Objectif: {objective[:80]}...")

        # 1. Parser les documents
        print("  [1/4] Parsing des documents...")
        all_text = ""
        for doc in documents:
            text = self.parse_document(
                file_path="",
                file_content=doc.get("content"),
                filename=doc.get("filename", "document.txt")
            )
            all_text += f"\n--- {doc.get('filename', 'Document')} ---\n{text}\n"

        if not all_text.strip():
            return {"error": "Aucun contenu extrait des documents."}

        # 2. Analyser et structurer
        slides = self.analyze_and_structure(all_text, target_audience, objective)

        if not slides:
            return {"error": "Impossible de structurer la presentation."}

        # 3. Generer les images
        slides = self.generate_images(slides)

        # 4. Construire le PPTX
        pptx_path = self.build_pptx(slides, title=objective[:50])

        return {
            "slides": slides,
            "pptx_path": pptx_path,
            "slide_count": len(slides),
            "generated_at": datetime.now().isoformat(),
        }
