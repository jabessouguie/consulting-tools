"""
Parser universel pour extraire du texte depuis différents formats de documents
Supporte : DOCX, PDF, TXT, MD
"""
import os
from pathlib import Path
from typing import Optional


class DocumentParser:
    """Extrait du texte depuis différents formats de documents"""

    @staticmethod
    def parse_file(file_path: str) -> Optional[str]:
        """
        Extrait le texte depuis un fichier (auto-détecte le format)

        Args:
            file_path: Chemin vers le fichier

        Returns:
            Texte extrait ou None si échec
        """
        file_path = Path(file_path)

        if not file_path.exists():
            print(f"❌ Fichier non trouvé : {file_path}")
            return None

        suffix = file_path.suffix.lower()

        if suffix in ['.txt', '.md', '.markdown']:
            return DocumentParser._parse_text(file_path)
        elif suffix == '.pdf':
            return DocumentParser._parse_pdf(file_path)
        elif suffix in ['.docx', '.doc']:
            return DocumentParser._parse_docx(file_path)
        else:
            print(f"⚠️  Format non supporté : {suffix}")
            return None

    @staticmethod
    def _parse_text(file_path: Path) -> Optional[str]:
        """Extrait le texte d'un fichier texte/markdown"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            print(f"✅ Texte extrait : {len(content)} caractères")
            return content
        except Exception as e:
            print(f"❌ Erreur lecture texte : {e}")
            return None

    @staticmethod
    def _parse_pdf(file_path: Path) -> Optional[str]:
        """Extrait le texte d'un fichier PDF"""
        try:
            # Essayer PyPDF2
            import PyPDF2

            text_content = []

            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                num_pages = len(pdf_reader.pages)

                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text:
                        text_content.append(text)

            content = '\n\n'.join(text_content)

            if content.strip():
                print(f"✅ PDF extrait : {len(content)} caractères ({num_pages} pages)")
                return content
            else:
                print("⚠️  Aucun texte extrait du PDF (possiblement un PDF image)")
                return None

        except ImportError:
            print("❌ PyPDF2 non installé. Installez avec: pip install PyPDF2")
            return None
        except Exception as e:
            print(f"❌ Erreur extraction PDF : {e}")
            return None

    @staticmethod
    def _parse_docx(file_path: Path) -> Optional[str]:
        """Extrait le texte d'un fichier DOCX"""
        try:
            # Essayer python-docx
            from docx import Document

            doc = Document(file_path)

            # Extraire les paragraphes
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

            # Extraire les tableaux
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        tables_text.append(row_text)

            # Combiner tout
            content_parts = paragraphs
            if tables_text:
                content_parts.append('\n--- TABLEAUX ---\n')
                content_parts.extend(tables_text)

            content = '\n\n'.join(content_parts)

            if content.strip():
                print(f"✅ DOCX extrait : {len(content)} caractères")
                return content
            else:
                print("⚠️  Aucun texte extrait du DOCX")
                return None

        except ImportError:
            print("❌ python-docx non installé. Installez avec: pip install python-docx")
            return None
        except Exception as e:
            print(f"❌ Erreur extraction DOCX : {e}")
            return None

    @staticmethod
    def is_format_supported(file_extension: str) -> bool:
        """
        Vérifie si un format de fichier est supporté

        Args:
            file_extension: Extension du fichier (avec ou sans point)

        Returns:
            True si supporté, False sinon
        """
        ext = file_extension.lower()
        if not ext.startswith('.'):
            ext = f'.{ext}'

        supported = ['.txt', '.md', '.markdown', '.pdf', '.docx', '.doc']
        return ext in supported

    @staticmethod
    def get_required_libraries() -> dict:
        """
        Retourne les bibliothèques requises pour chaque format

        Returns:
            Dict avec format → bibliothèque nécessaire
        """
        libraries = {
            'pdf': 'PyPDF2',
            'docx': 'python-docx',
            'txt': 'built-in',
            'md': 'built-in',
        }

        # Vérifier quelles bibliothèques sont installées
        status = {}
        for fmt, lib in libraries.items():
            if lib == 'built-in':
                status[fmt] = True
            else:
                try:
                    if fmt == 'pdf':
                        import PyPDF2
                        status[fmt] = True
                    elif fmt == 'docx':
                        import docx
                        status[fmt] = True
                except ImportError:
                    status[fmt] = False

        return status


# Instance globale
document_parser = DocumentParser()
