"""
Tests pour utils/word_exporter.py — export_to_word, _add_body_paragraph, _add_runs_with_bold
"""
import sys
from pathlib import Path
from unittest.mock import MagicMock, patch, call

import pytest


# ---------------------------------------------------------------------------
# ImportError path
# ---------------------------------------------------------------------------

class TestImportError:
    def test_raises_import_error_when_docx_not_installed(self, tmp_path):
        from utils.word_exporter import export_to_word

        with patch.dict("sys.modules", {"docx": None, "docx.shared": None}):
            with pytest.raises(ImportError, match="python-docx"):
                export_to_word({"title": "T", "sections": []}, str(tmp_path / "out.docx"))


# ---------------------------------------------------------------------------
# Helpers — _add_runs_with_bold
# ---------------------------------------------------------------------------

class TestAddRunsWithBold:
    def test_plain_text_not_bold(self):
        from utils.word_exporter import _add_runs_with_bold

        mock_para = MagicMock()
        runs = []

        def add_run(text):
            run = MagicMock()
            run.text = text
            runs.append(run)
            return run

        mock_para.add_run.side_effect = add_run
        _add_runs_with_bold(mock_para, "hello world")

        assert len(runs) == 1
        assert runs[0].bold is False

    def test_bold_text_is_bold(self):
        from utils.word_exporter import _add_runs_with_bold

        mock_para = MagicMock()
        runs = []

        def add_run(text):
            run = MagicMock()
            run.text = text
            runs.append(run)
            return run

        mock_para.add_run.side_effect = add_run
        _add_runs_with_bold(mock_para, "**important**")

        # "**important**".split("**") = ["", "important", ""]
        # Only "important" (i=1, odd) is bold; empty strings are skipped
        assert len(runs) == 1
        assert runs[0].bold is True

    def test_mixed_text_bold_and_normal(self):
        from utils.word_exporter import _add_runs_with_bold

        mock_para = MagicMock()
        runs = []

        def add_run(text):
            run = MagicMock()
            run.text = text
            runs.append(run)
            return run

        mock_para.add_run.side_effect = add_run
        _add_runs_with_bold(mock_para, "Start **bold** end")

        assert len(runs) == 3
        assert runs[0].bold is False   # "Start "
        assert runs[1].bold is True    # "bold"
        assert runs[2].bold is False   # " end"

    def test_empty_parts_are_skipped(self):
        from utils.word_exporter import _add_runs_with_bold

        mock_para = MagicMock()
        runs = []

        def add_run(text):
            run = MagicMock()
            runs.append(run)
            return run

        mock_para.add_run.side_effect = add_run
        _add_runs_with_bold(mock_para, "")
        assert len(runs) == 0

    def test_multiple_bold_segments(self):
        from utils.word_exporter import _add_runs_with_bold

        mock_para = MagicMock()
        runs = []

        def add_run(text):
            run = MagicMock()
            run.text = text
            runs.append(run)
            return run

        mock_para.add_run.side_effect = add_run
        # "**a** mid **b**" -> ["", "a", " mid ", "b", ""]
        # parts at odd indices (1, 3) are bold
        _add_runs_with_bold(mock_para, "**a** mid **b**")
        assert len(runs) == 3
        assert runs[0].bold is True   # "a"
        assert runs[1].bold is False  # " mid "
        assert runs[2].bold is True   # "b"

    def test_only_asterisks_no_content(self):
        from utils.word_exporter import _add_runs_with_bold

        mock_para = MagicMock()
        runs = []

        def add_run(text):
            run = MagicMock()
            runs.append(run)
            return run

        mock_para.add_run.side_effect = add_run
        # "****" -> ["", "", "", ""] — all empty, nothing added
        _add_runs_with_bold(mock_para, "****")
        assert len(runs) == 0


# ---------------------------------------------------------------------------
# Helpers — _add_body_paragraph
# ---------------------------------------------------------------------------

class TestAddBodyParagraph:
    def test_adds_plain_paragraph(self):
        from utils.word_exporter import _add_body_paragraph

        mock_doc = MagicMock()
        mock_para = MagicMock()
        mock_doc.add_paragraph.return_value = mock_para

        _add_body_paragraph(mock_doc, "Simple text")
        mock_doc.add_paragraph.assert_called()

    def test_bullet_item_uses_list_bullet_style(self):
        from utils.word_exporter import _add_body_paragraph

        mock_doc = MagicMock()
        mock_para = MagicMock()
        mock_doc.add_paragraph.return_value = mock_para

        _add_body_paragraph(mock_doc, "- first item")
        mock_doc.add_paragraph.assert_called_with(style="List Bullet")

    def test_bullet_with_bullet_unicode(self):
        from utils.word_exporter import _add_body_paragraph

        mock_doc = MagicMock()
        mock_para = MagicMock()
        mock_doc.add_paragraph.return_value = mock_para

        _add_body_paragraph(mock_doc, "• unicode bullet")
        mock_doc.add_paragraph.assert_called_with(style="List Bullet")

    def test_skips_empty_lines(self):
        from utils.word_exporter import _add_body_paragraph

        mock_doc = MagicMock()
        _add_body_paragraph(mock_doc, "\n\n\n")
        mock_doc.add_paragraph.assert_not_called()

    def test_multiline_text(self):
        from utils.word_exporter import _add_body_paragraph

        mock_doc = MagicMock()
        mock_para = MagicMock()
        mock_doc.add_paragraph.return_value = mock_para

        _add_body_paragraph(mock_doc, "Line 1\nLine 2\nLine 3")
        assert mock_doc.add_paragraph.call_count == 3

    def test_bullet_content_uses_add_runs_with_bold(self):
        """Bullet items strip '- ' prefix then pass remaining text to _add_runs_with_bold."""
        from utils.word_exporter import _add_body_paragraph

        mock_doc = MagicMock()
        mock_para = MagicMock()
        mock_doc.add_paragraph.return_value = mock_para

        _add_body_paragraph(mock_doc, "- **bold** item")
        mock_doc.add_paragraph.assert_called_with(style="List Bullet")
        # _add_runs_with_bold is called on the paragraph
        mock_para.add_run.assert_called()

    def test_regular_line_calls_add_runs(self):
        """Regular lines call add_paragraph() (no style arg) and then add_run."""
        from utils.word_exporter import _add_body_paragraph

        mock_doc = MagicMock()
        mock_para = MagicMock()
        mock_doc.add_paragraph.return_value = mock_para

        _add_body_paragraph(mock_doc, "Just text")
        mock_doc.add_paragraph.assert_called_once_with()
        mock_para.add_run.assert_called()


# ---------------------------------------------------------------------------
# export_to_word — integration tests (requires python-docx)
# ---------------------------------------------------------------------------

docx = pytest.importorskip("docx", reason="python-docx not installed")


class TestExportToWord:
    def test_creates_docx_file(self, tmp_path):
        from utils.word_exporter import export_to_word

        out = tmp_path / "output.docx"
        content = {"title": "My Doc", "sections": []}
        result = export_to_word(content, str(out))

        assert Path(result).exists()
        assert result.endswith(".docx")

    def test_returns_absolute_path(self, tmp_path):
        from utils.word_exporter import export_to_word

        out = tmp_path / "output.docx"
        result = export_to_word({"title": "T", "sections": []}, str(out))
        assert Path(result).is_absolute()

    def test_uses_title_parameter(self, tmp_path):
        from utils.word_exporter import export_to_word
        from docx import Document

        out = tmp_path / "out.docx"
        export_to_word({"title": "Original", "sections": []}, str(out), title="Override Title")

        doc = Document(str(out))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert any("Override Title" in h for h in headings)

    def test_uses_content_title_as_fallback(self, tmp_path):
        from utils.word_exporter import export_to_word
        from docx import Document

        out = tmp_path / "out.docx"
        export_to_word({"title": "Content Title", "sections": []}, str(out))

        doc = Document(str(out))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert any("Content Title" in h for h in headings)

    def test_default_title_when_no_title(self, tmp_path):
        from utils.word_exporter import export_to_word
        from docx import Document

        out = tmp_path / "out.docx"
        export_to_word({"sections": []}, str(out))

        doc = Document(str(out))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert any("Document" in h for h in headings)

    def test_creates_parent_directories(self, tmp_path):
        from utils.word_exporter import export_to_word

        out = tmp_path / "nested" / "deep" / "output.docx"
        result = export_to_word({"title": "T", "sections": []}, str(out))
        assert Path(result).exists()

    def test_sections_are_included(self, tmp_path):
        from utils.word_exporter import export_to_word
        from docx import Document

        out = tmp_path / "out.docx"
        content = {
            "title": "Report",
            "sections": [
                {"heading": "Introduction", "body": "Some intro text."},
                {"heading": "Conclusion", "body": "Final thoughts."},
            ],
        }
        export_to_word(content, str(out))

        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Introduction" in all_text
        assert "Conclusion" in all_text
        assert "Some intro text." in all_text

    def test_section_without_heading(self, tmp_path):
        from utils.word_exporter import export_to_word
        from docx import Document

        out = tmp_path / "out.docx"
        content = {
            "title": "T",
            "sections": [{"heading": "", "body": "No heading body."}],
        }
        export_to_word(content, str(out))
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "No heading body." in all_text

    def test_section_without_body(self, tmp_path):
        from utils.word_exporter import export_to_word
        from docx import Document

        out = tmp_path / "out.docx"
        content = {
            "title": "T",
            "sections": [{"heading": "Title Only", "body": ""}],
        }
        export_to_word(content, str(out))
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Title Only" in all_text

    def test_section_level_capped_at_3(self, tmp_path):
        from utils.word_exporter import export_to_word

        out = tmp_path / "out.docx"
        content = {
            "title": "T",
            "sections": [{"heading": "Deep", "body": "", "level": 10}],
        }
        # Should not raise
        export_to_word(content, str(out))
        assert Path(out).exists()

    def test_bold_text_in_body(self, tmp_path):
        from utils.word_exporter import export_to_word
        from docx import Document

        out = tmp_path / "out.docx"
        content = {
            "title": "T",
            "sections": [{"heading": "S", "body": "Normal **bold** text"}],
        }
        export_to_word(content, str(out))
        doc = Document(str(out))
        # Verify the paragraph with bold content exists
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "bold" in all_text

    def test_bullet_list_in_body(self, tmp_path):
        from utils.word_exporter import export_to_word
        from docx import Document

        out = tmp_path / "out.docx"
        content = {
            "title": "T",
            "sections": [{"heading": "S", "body": "- item one\n- item two"}],
        }
        export_to_word(content, str(out))
        doc = Document(str(out))
        styles = [p.style.name for p in doc.paragraphs]
        assert "List Bullet" in styles

    def test_unicode_bullet_in_body(self, tmp_path):
        from utils.word_exporter import export_to_word
        from docx import Document

        out = tmp_path / "out.docx"
        content = {
            "title": "T",
            "sections": [{"heading": "S", "body": "• first\n• second"}],
        }
        export_to_word(content, str(out))
        doc = Document(str(out))
        styles = [p.style.name for p in doc.paragraphs]
        assert "List Bullet" in styles

    def test_default_section_level_is_1(self, tmp_path):
        from utils.word_exporter import export_to_word
        from docx import Document

        out = tmp_path / "out.docx"
        content = {
            "title": "T",
            "sections": [{"heading": "Level Default", "body": ""}],
        }
        export_to_word(content, str(out))
        doc = Document(str(out))
        headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert any("Level Default" in h.text for h in headings)

    def test_multiple_sections_rendered(self, tmp_path):
        from utils.word_exporter import export_to_word
        from docx import Document

        out = tmp_path / "out.docx"
        sections = [
            {"heading": f"Section {i}", "body": f"Body {i}"}
            for i in range(5)
        ]
        export_to_word({"title": "Multi", "sections": sections}, str(out))
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        for i in range(5):
            assert f"Section {i}" in all_text
            assert f"Body {i}" in all_text
